import gradio as gr
import torch
import pytesseract
from PIL import Image
import numpy as np
from docx import Document
import tempfile
import os
import fitz  # PyMuPDF for PDF reading
import google.generativeai as genai
import psycopg2
import io
from urllib.parse import urlparse
from datetime import datetime
import time
import cv2
# === CONFIG ===
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
DATABASE_URL = os.environ.get("DATABASE_URL")
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("models/gemma-3n-e4b-it")

# === Database ===
def get_db_connection():
    return psycopg2.connect(DATABASE_URL, sslmode='require')

def create_table_if_not_exists():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS prescriptions (
            id SERIAL PRIMARY KEY,
            uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            filename TEXT,
            raw_text TEXT,
            structured_text TEXT
        )
    """)
    conn.commit()
    cursor.close()
    conn.close()

def save_to_neon_db(filename, raw_text, structured_text):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO prescriptions (filename, raw_text, structured_text)
            VALUES (%s, %s, %s)
        """, (filename, raw_text, structured_text))
        conn.commit()
        cursor.close()
        conn.close()
        return "✅ Saved to Neon DB."
    except Exception as e:
        return f"❌ Failed to save to DB: {e}"

# === OCR ===
def get_ocr_data(image):
    data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
    word_data = []
    for i in range(len(data['text'])):
        if int(data['conf'][i]) > 60 and data['text'][i].strip():
            word = data['text'][i]
            x, y, w, h = data['left'][i], data['top'][i], data['width'][i], data['height'][i]
            word_data.append({'text': word, 'box': [x, y, x + w, y + h], 'top': y, 'left': x})
    return word_data

def group_words_into_lines(word_data, line_threshold=15):
    word_data = sorted(word_data, key=lambda x: x['top'])
    lines, current_line, last_top = [], [], None
    for word in word_data:
        if last_top is None or abs(word['top'] - last_top) <= line_threshold:
            current_line.append(word)
        else:
            lines.append(sorted(current_line, key=lambda x: x['left']))
            current_line = [word]
        last_top = word['top']
    if current_line:
        lines.append(sorted(current_line, key=lambda x: x['left']))
    return lines

# === Main Extraction Logic ==
def process_prescription(file_path):
    """
    Faster, optimized processing for Gradio uploads.
    """
    start_total = time.time()

    if not file_path:
        return None, None, "❌ Please upload a file first."

    try:
        create_table_if_not_exists()
    except Exception as e:
        return None, None, f"❌ DB init failed: {e}"

    t0 = time.time()
    # Determine extension
    file_ext = os.path.splitext(file_path)[-1].lower().strip(".")
    # Copy uploaded file to temp (safe)
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}") as tmp_file:
        with open(file_path, "rb") as f:
            tmp_file.write(f.read())
        temp_path = tmp_file.name
    t1 = time.time()

    # Convert PDF first page to image with smaller dpi (150)
    if file_ext == "pdf":
        try:
            pdf = fitz.open(temp_path)
            pix = pdf[0].get_pixmap(dpi=150)   # <- lower DPI for speed
            image_bytes = pix.tobytes("png")
            image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        except Exception as e:
            return None, None, f"❌ PDF -> image failed: {e}"
    else:
        image = Image.open(temp_path).convert("RGB")
    t2 = time.time()

    # Convert PIL -> OpenCV (numpy) for fast preprocessing
    img_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

    # Resize to max width (preserve aspect ratio) to reduce OCR time
    max_width = 1500
    h, w = img_cv.shape[:2]
    if w > max_width:
        scale = max_width / w
        new_w = int(w * scale)
        new_h = int(h * scale)
        img_cv = cv2.resize(img_cv, (new_w, new_h), interpolation=cv2.INTER_AREA)

    # Grayscale + slight blur + adaptive threshold — faster and often more accurate
    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
    # optional denoising (only if images are noisy)
    gray = cv2.medianBlur(gray, 3)
    # adaptive threshold
    thresh = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY, 11, 2
    )

    # Convert back to PIL image for pytesseract
    preprocessed_pil = Image.fromarray(thresh)
    t3 = time.time()

    # OCR with optimized config
    tesseract_config = r'--oem 3 --psm 6 -l eng'  # tune lang if needed
    try:
        # Use image_to_data to keep word positions (like before)
        data = pytesseract.image_to_data(preprocessed_pil, output_type=pytesseract.Output.DICT, config=tesseract_config)
    except Exception as e:
        return None, None, f"❌ Tesseract OCR failed: {e}"
    t4 = time.time()

    # Build word_data (same structure as before)
    word_data = []
    n = len(data['text'])
    for i in range(n):
        conf = data.get('conf', [])[i]
        text = data.get('text', [])[i]
        if text and text.strip() and conf != '' and int(float(conf)) > 60:
            x, y, w_box, h_box = data['left'][i], data['top'][i], data['width'][i], data['height'][i]
            word_data.append({'text': text, 'box': [x, y, x + w_box, y + h_box], 'top': y, 'left': x})
    t5 = time.time()

    # Group words into lines (reuse your function)
    lines = group_words_into_lines(word_data)
    raw_text = "\n".join([" ".join([w['text'] for w in line]) for line in lines])
    t6 = time.time()

    # Save raw DOCX (fast)
    raw_docx_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    raw_doc = Document()
    raw_doc.add_heading('Raw Prescription Text', level=1)
    for line in lines:
        raw_doc.add_paragraph(" ".join([word['text'] for word in line]))
    raw_doc.save(raw_docx_path)
    t7 = time.time()

    # Gemini structuring — THIS can be the slowest step (network + model)
    # You can add a short prompt to ask for compact output to speed it slightly.
    prompt = f"""
Convert this raw OCR text into a structured prescription with:
Hospital Name, Patient Name, Age/Gender, Address, Date, Doctor Name & Degree,
Diagnosis, Medications (dosage & timing), Instructions, Follow-up Advice.

Raw OCR text:
\"\"\"{raw_text}\"\"\"

Return only the structured prescription.
"""
    try:
        t_gemini_start = time.time()
        response = model.generate_content(prompt)
        structured_text = response.text.strip()
        t_gemini_end = time.time()
    except Exception as e:
        return raw_docx_path, None, f"❌ Gemini structuring failed: {e}"

    # Save structured DOCX
    structured_docx_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    structured_doc = Document()
    structured_doc.add_heading('Structured Prescription', level=1)
    for paragraph in structured_text.split("\n\n"):
        structured_doc.add_paragraph(paragraph.strip())
    structured_doc.save(structured_docx_path)
    t8 = time.time()

    # Save to DB (keep it brief)
    try:
        db_status = save_to_neon_db(os.path.basename(file_path), raw_text, structured_text)
    except Exception as e:
        db_status = f"❌ DB save failed: {e}"

    total = time.time() - start_total

    # Make a small timing summary for debugging (remove in production)
    timing_summary = (
        f"timings (s): copy={t1-t0:.2f}, pdf_to_img={t2-t1:.2f}, preprocess={t3-t2:.2f}, "
        f"ocr={t4-t3:.2f}, build_words={t5-t4:.2f}, lines={t6-t5:.2f}, raw_save={t7-t6:.2f}, "
        f"gemini={t_gemini_end-t_gemini_start:.2f}, structured_save={t8-t_gemini_end:.2f}, total={total:.2f}"
    )

    return raw_docx_path, structured_docx_path, f"{db_status} | {timing_summary}"



# === Gradio UI ===
with gr.Blocks() as demo:
    gr.Markdown("## 🩺 Handwritten Prescription Structuring (Gradio)")

    with gr.Row():
        file_input = gr.File(label="Upload Prescription Image or PDF", file_types=[".jpg", ".jpeg", ".png", ".pdf"])
    
    start_btn = gr.Button("🚀 Start Extraction")

    with gr.Row():
        raw_download = gr.File(label="📄 Download Raw Prescription")
        structured_download = gr.File(label="📑 Download Structured Prescription")
    
    db_status_output = gr.Textbox(label="Database Status", interactive=False)

    start_btn.click(
        process_prescription,
        inputs=file_input,
        outputs=[raw_download, structured_download, db_status_output]
    )

if __name__ == "__main__":
    demo.launch(
        server_name="0.0.0.0",
        server_port=int(os.environ.get("PORT", 7860))
    )
